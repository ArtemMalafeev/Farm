from io import BytesIO

from docx import Document

from django.contrib import messages
from django.db.models import Sum
from django.http import Http404 , HttpResponse , JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views import View
from django.views.generic import DetailView, ListView

from fields.forms import FieldForm, JobForm
from fields.models import Crop, Field, Job


class FieldsView(ListView):
    """Список полей"""
    model = Field
    context_object_name = 'data'
    template_name = 'fields/field_list.html'

    def get_queryset(self):
        context = Crop.objects.raw('SELECT fc.id, ff.name, MAX(fs.name)'
                                   'FROM fields_crop fc '
                                   'JOIN fields_field ff ON fc.field_id = ff.id '
                                   'JOIN fields_season fs ON fc.season_id = fs.id '
                                   'GROUP BY ff.id')

        return context


class FieldDetail(DetailView):
    """Детальная информация о поле"""
    model = Field
    context_object_name = 'field'
    template_name = 'fields/field_detail.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['crops_list'] = self.object.crops.all().order_by('-sowing_date')  # Добавляем информацию по севам
        context['works_list'] = self.object.jobs.all().order_by('-start_job') # Добавляем информацию по работам

        return context


class FieldCreateView(View):
    """Создание нового поля"""
    def get(self, request):
        return render(request, 'fields/field_add.html', {'form': FieldForm()})

    def post(self, request):
        user = request.user
        form = FieldForm(request.POST)

        if form.is_valid():
            field = form.save(commit=False)
            field.user = user
            field.save()

            messages.add_message(request , messages.INFO , 'Поле успешно добавлено!')

            return redirect('/fields/add/')

        messages.add_message(request, messages.INFO, 'Поле с таким именем уже существует!')

        return render(request, 'fields/field_add.html', {'form': FieldForm()})


class FieldEditView(View):
    """Редактирование поля"""
    def get(self, request, slug):
        instance = get_object_or_404(Field, slug=slug)

        return render(request, 'fields/field_add.html',
                      {'form': FieldForm(instance=instance)})

    def post(self, request, slug):
        instance = get_object_or_404(Field, slug=slug)
        form = FieldForm(request.POST, instance=instance)

        if form.is_valid():
            form.save()

            return redirect('/fields/{}/'.format(slug))

        messages.add_message(request, messages.INFO, 'Данные неккоректны')

        return render(request, 'fields/field_add.html', {'form': FieldForm()})


class JobsDetail(DetailView):
    """ Вывод детальной информции о работе """
    model = Job
    context_object_name = 'job'
    template_name = 'fields/work.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['jobs'] = Job.objects.annotate(count_ha=Sum('fields__square_ha'),
                                    count_field=Sum('fields'))  # Добавляем информацию по севам

        return context


def job_add_view(request):
    """ Добавление новой работы """
    form = JobForm(request.POST or None, request.FILES or None)
    data = {}
    if request.is_ajax():
        if form.is_valid():
            print(form)
            form.save()
            data['status'] = 'ok'
            return JsonResponse(data)

    context = {
        'form': form,
    }

    return render(request, 'fields/work_add.html', context)


def export_work_docx(request, pk):
    """ Формирование документа в формате docx """
    work = Job.objects.get(pk=pk)

    document = Document()
    document.add_heading('Планирование работ', 0).bold = True

    docx_title = 'Job-{} {}.docx'.format(work.id, work.start_job.strftime("%d-%m-%Y"))

    document.add_heading('{}'.format(work.category), level=1)

    info = document.add_paragraph()
    info.add_run('    Сезон:                 ').bold = True
    info.add_run('{}'.format(work.season))

    info = document.add_paragraph()
    info.add_run('    Начало                 ').bold = True
    info.add_run('{}'.format(work.start_job.strftime("%d-%m-%Y")))

    info = document.add_paragraph()
    info.add_run('    Конец:                 ').bold = True
    info.add_run('{}'.format(work.end_job.strftime("%d-%m-%Y")))

    info = document.add_paragraph()
    info.add_run('    Продолжительность:                 ').bold = True
    info.add_run('{} дней'.format((work.end_job - work.start_job).days))

    document.add_heading('Список работников:', level=1)

    table = document.add_table(rows=1 , cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Имя'
    hdr_cells[2].text = 'Фамилия'
    count = 1

    for worker in work.workers.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        row_cells[1].text = worker.firstname
        row_cells[2].text = worker.lastname
        count += 1

    table.style = 'Colorful List'

    document.add_heading('Список полей:', level=1)

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Название'
    hdr_cells[2].text = 'Площадь м²'
    hdr_cells[3].text = 'Площадь га'
    count = 1

    for field in work.fields.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        row_cells[1].text = field.name
        row_cells[2].text = str(field.square)
        row_cells[3].text = str(field.square_ha)
        count += 1

    table.style = 'Colorful List'

    document.add_heading('Комментарии:', level=1)
    info = document.add_paragraph()

    if len(work.comment) != 0:
        info.add_run(work.comment)


    info = document.add_paragraph()
    info.add_run(' ')
    info = document.add_paragraph()
    info.add_run(' ')
    info = document.add_paragraph()
    info.add_run(' ')

    table = document.add_table(rows=1, cols=4)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Генеральный директор'
    row_cells[1].text = 'Малфеев А.Е.'
    row_cells[2].text = '_______________'
    row_cells[3].text = 'Утверждено'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Главный агроном'
    row_cells[1].text = 'Косыгина А.В.'
    row_cells[2].text = '_______________'
    row_cells[3].text = 'Утверждено'

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)

    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length

    return response
